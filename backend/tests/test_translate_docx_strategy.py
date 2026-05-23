"""Tests for DOCX translation strategy behavior."""

from pathlib import Path

from docx import Document

from app.translate.domain.dto import SegmentRecord
from app.translate.service.strategies.docx_strategy import DocxTranslateStrategy


def test_docx_translates_paragraphs_and_table_cells(tmp_path: Path) -> None:
    """Translate DOCX paragraphs and table cells without dropping run style."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Hello")
    run.bold = True
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell text"
    doc.save(source)

    strategy = DocxTranslateStrategy()
    drafts = strategy.extract(source)
    records = [
        SegmentRecord(
            seq=d.seq,
            source_text=d.source_text,
            translated_text=f"译:{d.source_text}",
            anchor_json=d.anchor_json,
        )
        for d in drafts
    ]
    strategy.assemble(records, source, output)

    out = Document(output)
    assert out.paragraphs[0].text == "译:Hello"
    assert out.paragraphs[0].runs[0].bold is True
    assert out.tables[0].cell(0, 0).text == "译:Cell text"


def test_docx_table_cell_writeback_clears_extra_paragraphs(tmp_path: Path) -> None:
    """Translate a multi-paragraph table cell without leaving stale source text."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "First"
    cell.add_paragraph("Second")
    original_cell_text = cell.text
    doc.save(source)

    assert original_cell_text == "First\nSecond"

    strategy = DocxTranslateStrategy()
    draft = next(
        d
        for d in strategy.extract(source)
        if d.anchor_json and d.anchor_json["kind"] == "table_cell"
    )
    assert draft.source_text == original_cell_text
    record = SegmentRecord(
        seq=draft.seq,
        source_text=draft.source_text,
        translated_text=f"译:{draft.source_text}",
        anchor_json=draft.anchor_json,
    )
    strategy.assemble([record], source, output)

    out_cell_text = Document(output).tables[0].cell(0, 0).text
    assert out_cell_text.rstrip("\n") == f"译:{original_cell_text}"
    assert out_cell_text.count("First") == 1
    assert out_cell_text.count("Second") == 1

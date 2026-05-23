"""DOCX writer preserving paragraph and table-cell structure where possible."""

from __future__ import annotations

from docx import Document
from docx.text.paragraph import Paragraph

from app.layout.writers.base import WriteContext


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace text while preserving the first run's basic style."""

    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


class DocxWriter:
    """Write translated segments into DOCX paragraphs and table cells."""

    def write(self, context: WriteContext) -> None:
        """Create the translated DOCX at ``context.out_path``."""

        doc = Document(context.source_path)
        for seg in context.segments:
            anchor = seg.anchor_json or {}
            text = (
                seg.source_text
                if anchor.get("skip_translate")
                else seg.translated_text or seg.source_text
            )
            kind = anchor.get("kind")
            if kind == "paragraph":
                idx = int(anchor.get("index", 0))
                if 0 <= idx < len(doc.paragraphs):
                    _replace_paragraph_text(doc.paragraphs[idx], text)
            elif kind == "table_cell":
                t_idx = int(anchor.get("table", 0))
                r_idx = int(anchor.get("row", 0))
                c_idx = int(anchor.get("col", 0))
                if 0 <= t_idx < len(doc.tables):
                    table = doc.tables[t_idx]
                    row_cells = []
                    if 0 <= r_idx < len(table.rows):
                        row_cells = table.rows[r_idx].cells
                    if 0 <= c_idx < len(row_cells):
                        cell = row_cells[c_idx]
                        if cell.paragraphs:
                            _replace_paragraph_text(cell.paragraphs[0], text)
                        else:
                            cell.text = text
        doc.save(context.out_path)

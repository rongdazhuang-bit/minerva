"""Word docx document translation strategy."""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import ClassVar

from docx import Document

from app.layout.writers.base import WriteContext
from app.layout.writers.docx_writer import DocxWriter
from app.translate.domain.dto import SegmentDraft, SegmentRecord
from app.translate.service.strategies.base import DocTranslateFormatStrategy


class DocxTranslateStrategy(DocTranslateFormatStrategy):
    """Translate ``.docx`` paragraphs and table cells preserving structure."""

    extensions: ClassVar[frozenset[str]] = frozenset({"docx"})

    def extract(
        self,
        local_path: Path,
        *,
        ocr_file_id: uuid.UUID | None = None,
        ocr_pages: list[tuple[int, str]] | None = None,
        layout_document=None,
    ) -> list[SegmentDraft]:
        """Extract non-empty DOCX paragraphs and table cells for translation."""

        doc = Document(local_path)
        drafts: list[SegmentDraft] = []
        seq = 0
        for p_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            drafts.append(
                SegmentDraft(
                    seq=seq,
                    source_text=text,
                    anchor_json={
                        "kind": "paragraph",
                        "index": p_idx,
                        "block_key": f"p0.p{p_idx}",
                        "label": "text",
                        "page_index": 0,
                        "overflow_policy": "shrink",
                    },
                )
            )
            seq += 1
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    drafts.append(
                        SegmentDraft(
                            seq=seq,
                            source_text=text,
                            anchor_json={
                                "kind": "table_cell",
                                "table": t_idx,
                                "row": r_idx,
                                "col": c_idx,
                                "block_key": f"p0.t{t_idx}.r{r_idx}.c{c_idx}",
                                "label": "table_cell",
                                "page_index": 0,
                                "overflow_policy": "expand",
                            },
                        )
                    )
                    seq += 1
        return drafts

    def assemble(
        self,
        segments: list[SegmentRecord],
        source_path: Path,
        out_path: Path,
    ) -> None:
        """Assemble translated DOCX content while preserving editable structure."""

        DocxWriter().write(
            WriteContext(source_path=source_path, out_path=out_path, segments=segments)
        )
